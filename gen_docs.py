"""
运行方式：
  cd /mnt/d/DevSoftware/code/forum
  source venv/bin/activate
  python3 gen_docs.py
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

# ============================================
# 工具函数
# ============================================
def shade_cell(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def make_header_row(table, headers, color='2F5496'):
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        shade_cell(hdr[i], color)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

# ============================================
# 生成测试计划文档
# ============================================
def gen_test_plan():
    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)

    title = doc.add_heading('Forum API 测试计划', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ver.add_run('版本 v1.0').font.size = Pt(10)

    doc.add_paragraph()

    # 1. 项目概述
    doc.add_heading('1. 项目概述', 1)
    doc.add_paragraph(
        'Forum API 是一个基于 FastAPI + MySQL + Redis 的社交论坛后端服务，'
        '提供用户认证、帖子管理、点赞互动等核心功能。'
        '本测试计划旨在通过系统化的测试策略，确保 API 的功能完整性、性能可靠性和安全性。'
    )

    # 2. 测试范围
    doc.add_heading('2. 测试范围', 1)
    doc.add_paragraph('【测试内容】用户注册/登录/登出、帖子发布/查看、点赞/取消点赞、JWT Token 黑名单验证、Redis 速率限制', style='List Bullet')
    doc.add_paragraph('【不测试】UI 界面、前端应用、第三方依赖库内部实现', style='List Bullet')
    doc.add_paragraph('【测试深度】单元测试 + 集成测试 + 性能测试（三层测试金字塔全覆盖）', style='List Bullet')

    # 3. 测试策略
    doc.add_heading('3. 测试策略', 1)
    doc.add_heading('3.1 分层测试（Testing Pyramid）', 2)
    t1 = doc.add_table(rows=4, cols=3)
    t1.style = 'Light Grid Accent 1'
    make_header_row(t1, ['测试层级', '覆盖范围', '工具 & 指标'])
    for idx, row in enumerate([
        ['单元测试', 'Mocking 隔离业务逻辑（密码哈希、JWT 验证）', 'Pytest + Mock，覆盖率 >90%'],
        ['集成测试', '真实数据库 + Redis，端到端 API 流程验证', 'Pytest + TestClient，27+ DDT 参数化用例'],
        ['性能测试', '并发压力、吞吐量、响应时间、连接池极限', 'Locust FastHttpUser，100 并发，60 秒'],
    ]):
        cells = t1.rows[idx + 1].cells
        for i, text in enumerate(row):
            cells[i].text = text

    doc.add_paragraph()
    doc.add_heading('3.2 测试设计方法', 2)
    methods = [
        '等价类划分（Equivalence Partitioning）：用户名长度分类（无效区间：<3 或 >20；有效区间：3-20），密码复杂度分类',
        '边界值分析（Boundary Value Analysis）：用户名长度（2/3/20/21），帖子标题长度（0/1/100/101字符）',
        '错误猜测法（Error Guessing）：空值、纯空格、SQL注入字符、超长字符串、无效 Token 格式',
        '场景法（Scenario Testing）：多用户交互链路（用户 A 发帖 → 用户 B 点赞 → B 取消点赞 → 验证状态一致性）',
    ]
    for m in methods:
        doc.add_paragraph(m, style='List Bullet')

    # 4. 测试环境
    doc.add_heading('4. 测试环境', 1)
    t2 = doc.add_table(rows=6, cols=3)
    t2.style = 'Light Grid Accent 1'
    make_header_row(t2, ['组件', '版本', '部署方式'])
    for idx, row in enumerate([
        ['Python', '3.11+', 'WSL 本地虚拟环境'],
        ['FastAPI', 'Latest', 'Uvicorn ASGI 服务器'],
        ['MySQL', '8.0', 'Docker 容器'],
        ['Redis', '7.0-alpine', 'Docker 容器'],
        ['GitHub Actions', 'ubuntu-latest', 'CI/CD 自动化流水线'],
    ]):
        cells = t2.rows[idx + 1].cells
        for i, text in enumerate(row):
            cells[i].text = text

    # 5. 进入/退出准则
    doc.add_heading('5. 进入 / 退出准则', 1)
    t3 = doc.add_table(rows=3, cols=2)
    t3.style = 'Light Grid Accent 1'
    make_header_row(t3, ['阶段', '准则'])
    criteria = [
        ['进入准则', '✓ 代码合并至 main 分支\n✓ 数据库迁移已执行（alembic upgrade head）\n✓ 所有依赖已安装'],
        ['退出准则', '✓ 代码覆盖率 ≥ 90%\n✓ 全部用例通过（0 failures）\n✓ 性能测试无雪崩现象\n✓ 无 P1 级别未修复缺陷'],
    ]
    for i, (stage, text) in enumerate(criteria):
        cells = t3.rows[i + 1].cells
        cells[0].text = stage
        cells[1].text = text

    # 6. 风险分析
    doc.add_heading('6. 风险分析与缓解措施', 1)
    t4 = doc.add_table(rows=4, cols=3)
    t4.style = 'Light Grid Accent 1'
    make_header_row(t4, ['风险项', '影响等级', '缓解措施'])
    risks = [
        ['MySQL 连接池在高并发下耗尽', '高', '压测中精确捕获 OperationalError，验证 max_connections 配置'],
        ['JWT Token 过期后未及时黑名单', '中', '集成测试覆盖登出后 Token 重放攻击场景'],
        ['Redis 防刷在并发下失效', '中', '边界测试：精确验证 10 秒 3 次限制逻辑'],
    ]
    for idx, row in enumerate(risks):
        cells = t4.rows[idx + 1].cells
        for i, text in enumerate(row):
            cells[i].text = text

    # 7. 交付物
    doc.add_heading('7. 预期交付物', 1)
    deliverables = [
        'Pytest 单元/集成测试代码（tests/ 目录，覆盖率报告）',
        'Allure 可视化测试报告（自动发布至 GitHub Pages）',
        'Locust 性能测试报告（RPS、响应时间、雪崩分析）',
        '本测试计划文档 + 测试用例矩阵',
    ]
    for d in deliverables:
        doc.add_paragraph(d, style='List Bullet')

    doc.add_paragraph()
    footer = doc.add_paragraph('--- End of Test Plan ---')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].italic = True

    doc.save('测试计划_v1.0.docx')
    print('✅ 测试计划_v1.0.docx 生成成功')


# ============================================
# 生成测试用例矩阵
# ============================================
def gen_test_matrix():
    wb = Workbook()
    sheet = wb.active
    sheet.title = '测试用例矩阵'

    col_widths = [8, 10, 14, 40, 16, 16, 10, 10]
    for i, w in enumerate(col_widths, 1):
        sheet.column_dimensions[chr(64 + i)].width = w
    sheet.row_dimensions[1].height = 30

    headers = ['用例ID', '模块', '功能点', '测试用例描述', '前置条件', '设计方法', '优先级', '预期结果']
    hfill = PatternFill(start_color='2F5496', end_color='2F5496', fill_type='solid')
    hfont = Font(bold=True, color='FFFFFF', size=11)
    bd = Border(left=Side(style='thin'), right=Side(style='thin'),
                top=Side(style='thin'), bottom=Side(style='thin'))
    center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    wrap = Alignment(vertical='center', wrap_text=True)

    for col, h in enumerate(headers, 1):
        cell = sheet.cell(row=1, column=col)
        cell.value = h
        cell.fill = hfill
        cell.font = hfont
        cell.alignment = center
        cell.border = bd

    cases = [
        # Auth 模块
        ['A001', '认证', '用户注册', '正常用户名（字母数字组合）成功注册', '服务启动', '等价类划分', 'P1', '返回 code:200'],
        ['A002', '认证', '用户注册', '用户名最小边界（3字符）成功注册', '服务启动', '边界值分析', 'P1', '返回 code:200'],
        ['A003', '认证', '用户注册', '用户名最大边界（20字符）成功注册', '服务启动', '边界值分析', 'P1', '返回 code:200'],
        ['A004', '认证', '用户注册', '用户名过短（2字符）注册失败', '服务启动', '边界值分析', 'P1', '返回 code:400'],
        ['A005', '认证', '用户注册', '用户名过长（21字符）注册失败', '服务启动', '边界值分析', 'P1', '返回 code:400'],
        ['A006', '认证', '用户注册', '用户名含特殊字符（@#$）注册失败', '服务启动', '错误猜测', 'P1', '返回 code:400'],
        ['A007', '认证', '用户注册', '密码含大小写+数字+特殊字符注册成功', '服务启动', '等价类划分', 'P1', '返回 code:200'],
        ['A008', '认证', '用户注册', '密码最小边界（8字符）注册成功', '服务启动', '边界值分析', 'P1', '返回 code:200'],
        ['A009', '认证', '用户注册', '密码最大边界（16字符）注册成功', '服务启动', '边界值分析', 'P1', '返回 code:200'],
        ['A010', '认证', '用户注册', '密码过短（7字符）注册失败', '服务启动', '边界值分析', 'P1', '返回 code:400'],
        ['A011', '认证', '用户注册', '密码缺少大写字母注册失败', '服务启动', '错误猜测', 'P1', '返回 code:400'],
        ['A012', '认证', '用户注册', '密码缺少数字注册失败', '服务启动', '错误猜测', 'P1', '返回 code:400'],
        ['A013', '认证', '用户注册', '有效邮箱格式注册成功', '服务启动', '等价类划分', 'P1', '返回 code:200'],
        ['A014', '认证', '用户注册', '邮箱缺少@符号注册失败', '服务启动', '错误猜测', 'P1', '返回 code:400'],
        ['A015', '认证', '用户注册', '重复用户名注册失败', '已存在用户', '错误猜测', 'P1', '返回 code:409'],
        ['A016', '认证', '用户登录', '正确用户名+密码登录成功并返回 Token', '已注册用户', '等价类划分', 'P1', '返回 JWT Token'],
        ['A017', '认证', '用户登录', '错误密码登录失败', '已注册用户', '错误猜测', 'P1', '返回 code:401'],
        ['A018', '认证', '用户登录', '不存在用户登录失败', '服务启动', '错误猜测', 'P1', '返回 code:401'],
        ['A019', '认证', '用户登出', '持有有效 Token 登出成功', '已登录', '等价类划分', 'P1', '返回 code:200'],
        ['A020', '认证', '用户登出', '登出后 Token 写入 Redis 黑名单', '已登出', '场景法', 'P1', 'Redis 中存在该 Token'],
        ['A021', '认证', '用户登出', '黑名单 Token 访问受保护接口被拒绝', '已登出', '场景法', 'P1', '返回 code:401'],
        # Posts 模块
        ['P001', '帖子', '发布帖子', '标题最小边界（1字符）发布成功', '已登录', '边界值分析', 'P1', '返回 code:200 + post_id'],
        ['P002', '帖子', '发布帖子', '标题最大边界（100字符）发布成功', '已登录', '边界值分析', 'P1', '返回 code:200 + post_id'],
        ['P003', '帖子', '发布帖子', '标题过长（101字符）发布失败', '已登录', '边界值分析', 'P1', '返回 code:400'],
        ['P004', '帖子', '发布帖子', '正文为空发布失败', '已登录', '错误猜测', 'P1', '返回 code:400'],
        ['P005', '帖子', '发布帖子', '正文仅空格发布失败', '已登录', '错误猜测', 'P1', '返回 code:400'],
        ['P006', '帖子', '发布帖子', '未携带 Token 发布失败', '未登录', '错误猜测', 'P1', '返回 code:401'],
        ['P007', '帖子', '速率限制', '10秒内发3帖成功（边界内）', '已登录', '边界值分析', 'P1', '均返回 code:200'],
        ['P008', '帖子', '速率限制', '10秒内发第4帖被限流', '已登录', '边界值分析', 'P1', '返回 code:429'],
        ['P009', '帖子', '获取列表', '默认分页（limit=10, offset=0）', '已有帖子', '等价类划分', 'P1', '返回帖子数组'],
        ['P010', '帖子', '获取列表', '自定义分页参数', '已有帖子', '等价类划分', 'P2', '返回对应数量帖子'],
        ['P011', '帖子', '点赞', '首次点赞成功，likes_count +1', '已登录', '场景法', 'P1', 'likes_count 递增'],
        ['P012', '帖子', '点赞', '再次点赞取消，likes_count -1', '已点赞', '场景法', 'P1', 'likes_count 递减'],
        ['P013', '帖子', '点赞', '多用户交互后 likes_count 数值一致', '多用户', '场景法', 'P1', '数据库与返回值一致'],
    ]

    p1_fill = PatternFill(start_color='FFE699', end_color='FFE699', fill_type='solid')

    for row_idx, case in enumerate(cases, 2):
        for col_idx, value in enumerate(case, 1):
            cell = sheet.cell(row=row_idx, column=col_idx)
            cell.value = value
            cell.border = bd
            if col_idx in [1, 7]:
                cell.alignment = center
                if col_idx == 7 and value == 'P1':
                    cell.fill = p1_fill
            else:
                cell.alignment = wrap

    sheet.freeze_panes = 'A2'
    sheet.auto_filter.ref = f'A1:H{len(cases) + 1}'

    # 统计页
    summary = wb.create_sheet('统计')
    summary.column_dimensions['A'].width = 16
    summary.column_dimensions['B'].width = 10
    summary.column_dimensions['C'].width = 10
    summary.column_dimensions['D'].width = 10
    summary.column_dimensions['E'].width = 12

    summary['A1'] = '测试用例统计'
    summary['A1'].font = Font(bold=True, size=14)

    for col, text in enumerate(['模块', '总数', 'P1', 'P2', '设计方法覆盖'], 1):
        cell = summary.cell(row=3, column=col)
        cell.value = text
        cell.fill = hfill
        cell.font = hfont
        cell.border = bd
        cell.alignment = center

    stats = [
        ['认证（Auth）', 21, 21, 0, '等价类、边界值、错误猜测、场景法'],
        ['帖子（Posts）', 13, 12, 1, '边界值、错误猜测、场景法'],
        ['合计', 34, 33, 1, '4 种方法全覆盖'],
    ]
    for row_idx, row in enumerate(stats, 4):
        for col_idx, value in enumerate(row, 1):
            cell = summary.cell(row=row_idx, column=col_idx)
            cell.value = value
            cell.border = bd
            cell.alignment = center if col_idx != 5 else wrap
            if row_idx == 6:
                cell.font = Font(bold=True)

    wb.save('测试用例矩阵_v1.0.xlsx')
    print('✅ 测试用例矩阵_v1.0.xlsx 生成成功')


if __name__ == '__main__':
    gen_test_plan()
    gen_test_matrix()
    print('\n🎉 全部文档生成完毕！文件位置：当前目录')
